# -*- coding: utf-8 -*-
"""更新桌面「技術搞2.docx」：內嵌架構圖換成最新 PNG，並將正文／表格轉成臺灣繁體。

依賴：pip install python-docx zhconv
圖檔：先執行 python docs/figures/render_architecture_png.py

若 Word 正在開啟該檔，會寫入 %TEMP%\\技術搞2_已更新.docx，請關閉 Word 後手動覆蓋。
"""
from __future__ import annotations

import hashlib
import os
import shutil
import subprocess
import sys
import tempfile
import zipfile
from pathlib import Path

try:
    import zhconv
except ImportError:
    print("請先執行: pip install zhconv", file=sys.stderr)
    raise SystemExit(1)

from docx import Document

def _find_desktop_tech_docx() -> Path | None:
    """Windows 下 Path.glob 對部分中文路徑不可靠時，改由 PowerShell 回傳 UTF-8 路徑。"""
    home = Path.home()
    for p in home.glob("OneDrive*/Desktop/\u6280\u8853\u641e2.docx"):
        if p.is_file():
            try:
                if p.stat().st_size > 0:
                    return p
            except OSError:
                continue
    ps = (
        "Get-ChildItem -LiteralPath (Join-Path $env:USERPROFILE (Get-ChildItem "
        "$env:USERPROFILE\\OneDrive* | Select-Object -First 1 | ForEach-Object { $_.Name })\\Desktop "
        "-Filter '*.docx' -ErrorAction SilentlyContinue | "
        "Where-Object { $_.Name -eq \"\u6280\u8853\u641e2.docx\" } | "
        "Select-Object -ExpandProperty FullName"
    )
    try:
        cp = subprocess.run(
            ["powershell", "-NoProfile", "-Command", ps],
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="replace",
            timeout=30,
        )
        out = (cp.stdout or "").strip()
        if out and Path(out).is_file():
            return Path(out)
    except (OSError, subprocess.TimeoutExpired):
        pass
    return None


def _to_tw(text: str) -> str:
    if not text or not text.strip():
        return text
    return zhconv.convert(text, "zh-tw")


def _embedded_images_match(docx: Path, arch: Path, seq: Path) -> bool:
    """確認 docx 內 word/media/image1/2.png 與專案產出之 PNG 位元組一致。"""
    want1 = hashlib.sha256(arch.read_bytes()).digest()
    want2 = hashlib.sha256(seq.read_bytes()).digest()
    with zipfile.ZipFile(docx, "r") as z:
        got1 = hashlib.sha256(z.read("word/media/image1.png")).digest()
        got2 = hashlib.sha256(z.read("word/media/image2.png")).digest()
    return got1 == want1 and got2 == want2


def _apply_tw_doc(doc: Document) -> None:
    for p in doc.paragraphs:
        if p.text.strip():
            p.text = _to_tw(p.text)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                if cell.text.strip():
                    cell.text = _to_tw(cell.text)


def main() -> int:
    root = Path(__file__).resolve().parents[1]
    arch = root / "docs" / "figures" / "architecture_overview.png"
    seq = root / "docs" / "figures" / "dataflow_sequence.png"
    if not arch.is_file() or not seq.is_file():
        print("找不到 PNG，請先執行: python docs/figures/render_architecture_png.py", file=sys.stderr)
        return 1

    target = _find_desktop_tech_docx()
    if target is None:
        print("找不到桌面 技術搞2.docx", file=sys.stderr)
        return 1

    # 先複製到僅 ASCII 路徑再開啟，避免部分環境下 python-docx 對中文路徑解包失敗
    work_dir = Path(tempfile.mkdtemp(prefix="docx_tw_"))
    src_ascii = work_dir / "source.docx"
    shutil.copy2(target, src_ascii)

    tmp_doc = work_dir / "patched.docx"
    doc = Document(str(src_ascii))
    _apply_tw_doc(doc)
    doc.save(str(tmp_doc))

    b1 = arch.read_bytes()
    b2 = seq.read_bytes()
    out_buf = Path(tempfile.mktemp(suffix=".docx"))
    with zipfile.ZipFile(tmp_doc, "r") as zin:
        with zipfile.ZipFile(out_buf, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "word/media/image1.png":
                    data = b1
                elif item.filename == "word/media/image2.png":
                    data = b2
                zout.writestr(item, data)

    try:
        shutil.copy2(out_buf, target)
        if not _embedded_images_match(target, arch, seq):
            print(
                "[錯誤] 寫入後內嵌圖雜湊仍與 docs/figures 不符（常見：OneDrive 還原舊版、或檔案仍被鎖）。",
                file=sys.stderr,
            )
            print("       請暫停編輯該檔、關閉 Word、待 OneDrive 同步完成後再執行本腳本。", file=sys.stderr)
            return 3
        print("已更新:", target)
        print("[確認] 內嵌圖與 architecture_overview.png / dataflow_sequence.png 一致。")
    except OSError as e:
        fallback = Path(os.environ["TEMP"]) / "tech2_docx_updated.docx"
        shutil.copy2(out_buf, fallback)
        print("無法直接覆寫（可能 Word 仍開著原檔）。已寫入:", fallback)
        print("錯誤:", e)
        return 2
    finally:
        shutil.rmtree(work_dir, ignore_errors=True)
        out_buf.unlink(missing_ok=True)

    return 0


if __name__ == "__main__":
    raise SystemExit(main())

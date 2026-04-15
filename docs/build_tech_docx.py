# 產生簡報用技術說明 Word（.docx）
# 執行：python docs/build_tech_docx.py
from __future__ import annotations

import pathlib
from datetime import date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

HERE = pathlib.Path(__file__).resolve().parent
OUT = HERE / "child_tracker_kg_技術說明.docx"

ROWS = [
    ("語言", "Python 3", "主要實作、CLI、影像與圖譜管線"),
    ("影像／影片", "OpenCV（opencv-python）", "讀寫影片、畫框與標籤、部分追蹤輔助"),
    ("人體偵測／追蹤", "Ultralytics YOLOv8（ultralytics）", "人體框、跨幀軌道；程式載入 yolov8{n,s,m,l,x}.pt"),
    ("臉部辨識", "face-recognition（底層 dlib）", "128 維臉部編碼、與註冊庫距離比對、門檻與 margin 策略"),
    ("數值運算", "NumPy、Pillow", "陣列運算、影像載入；可選 SciPy（最佳指派等）"),
    ("註冊與資料", "face_registry.py、data/registered/、data/face_registry.json", "幼兒照片與模版註冊，非關聯式資料庫"),
    ("追蹤實作", "yolo_tracker.py、video_tracker.py", "YOLO＋臉部；或純臉部（fallback、--no-yolo）"),
    ("互動紀錄", "data/interactions.json", "同框、距離靠近；YOLO 回溯模式下可含 interaction_timeline（JSON）"),
    ("圖結構", "NetworkX", "建「幼兒—朋友」加權圖、匯出 knowledge_graph.json"),
    ("視覺化", "pyvis、vis-network（產出 HTML）", "可互動拖曳、縮放之關係圖"),
    ("圖上時間軸", "Chart.js（CDN 注入 HTML）", "節點 hover 顯示互動折線圖（需時間軸資料）"),
    ("個性標籤（代理）", "personality.py", "依全班互動分位數貼標（proxy，非正式量表）"),
    ("CLI", "main.py（register / process / build-graph / run-all 等）", "單一入口；YOLO 預設可經 subprocess 呼叫 run_yolo_isolated.py"),
    ("執行環境穩定", "runtime_hardening.py、環境變數與執行緒限制", "降低 Windows 上 PyTorch／dlib／OpenMP 混用時不穩定"),
    ("部署型態", "離線本機、靜態 HTML", "無後端伺服器、無 SQL 資料庫；以檔案系統保存輸出"),
]


def _set_cell_text(cell, text: str):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.font.size = Pt(10)


def main():
    doc = Document()

    t = doc.add_heading("child_tracker_kg 技術說明（簡報用稿）", level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.add_run(
        "本文件依專案程式與設定整理，供簡報／報告直接摘用。預設輸出影片路徑：data/output/output.mp4（可用 --output 指定）。"
    )

    doc.add_paragraph(f"產生日期：{date.today().isoformat()}")

    doc.add_heading("技術棧（Tech Stack）", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    _set_cell_text(hdr[0], "類別")
    _set_cell_text(hdr[1], "技術／工具")
    _set_cell_text(hdr[2], "用途（本專案中的角色）")
    for c0, c1, c2 in ROWS:
        row = table.add_row().cells
        _set_cell_text(row[0], c0)
        _set_cell_text(row[1], c1)
        _set_cell_text(row[2], c2)

    doc.add_heading("1. 技術架構總覽", level=1)
    for line in [
        "專案名稱：child_tracker_kg（幼兒辨識、追蹤與關係／知識圖譜管線）。",
        "架構型態：本機離線管線；輸入為註冊照與影片，輸出為標註影片、JSON 互動、互動式關係圖 HTML。",
        "區塊劃分建議：輸入層 → 本機執行（main.py、可選 YOLO 子行程）→ 擇一追蹤核心（yolo_tracker 或 video_tracker）→ 寫入 interactions.json → data/ 檔案儲存 → build-graph（NetworkX、personality、pyvis）→ 瀏覽器開啟靜態 HTML（vis-network、Chart.js）。",
        "資料儲存：以 data/ 目錄為主（影像、JSON、HTML），不使用傳統資料庫。",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    doc.add_heading("2. 關鍵技術與理由", level=1)
    for line in [
        "YOLO 人體框＋臉部比對：多幼兒同畫面時，用人體軌道穩定「誰在哪」，再掛名；較單純全畫面臉追蹤利於複雜教室場景。",
        "run_yolo_isolated.py 子行程：與主程式內已載入之 dlib 等分離，減少 PyTorch／OpenMP 衝突與無訊息崩潰（搭配 runtime_hardening）。",
        "距離與 margin、累積分數、換名冷卻、強匹配覆蓋：降低相似臉誤掛與單幀雜訊；參數集中於 config.py。",
        "教室預設：RECOGNITION_PRESET（如 classroom_masked）可連動較符合實場的門檻與追蹤模式。",
        "NetworkX＋pyvis：快速由互動 JSON 建圖並產免伺服器之互動網頁。",
        "Chart.js 時間軸：將 interaction_timeline（JSON）在節點 hover 時視覺化；並非獨立輸出 interaction_timeline.png。",
        "個性模組：僅能陳述為互動行為之代理指標，避免宣稱等同正式氣質／人格量表。",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    doc.add_heading("3. 技術實作歷程", level=1)
    for line in [
        "典型流程：register 寫入註冊 → process 跑影片（YOLO 或純臉部）→ 寫入 interactions.json → build-graph 讀 JSON 建 knowledge_graph.json 並輸出 relationship_graph.html 等。",
        "時序要點：建圖階段為 interactions.json → 建圖／繪圖邏輯（讀檔後建 NetworkX，再 pyvis／注入 Chart.js）。",
        "核心檔案（模組—職責）：main.py、face_registry.py、yolo_tracker.py／video_tracker.py、knowledge_graph.py、relationship_graph.py、personality.py、config.py。",
        "延伸功能（選寫）：build-graph --ego 個人子圖；extract-faces；audit_registry 註冊照稽核。",
        "技術挑戰（實作對應）：Windows 混合執行緒／兩套數值庫衝突；相似臉與並排幼兒；YOLO 失敗時 fallback 至純臉部（折線時間軸可能為空，需如實說明）。",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    doc.add_heading("4. 技術特色與優勢", level=1)
    for line in [
        "相較純人工計數同框：可對整部影片批次統計，並在回溯模式下提供時間桶級互動結構（供圖上折線圖）。",
        "相較需架站之 Web 系統：產出靜態 HTML，利於示範、離線簡報。",
        "模組化：追蹤、註冊、建圖、視覺化分界清楚，方便換場景時只調 config 與預設。",
        "成果驗證建議（自填數據）：處理時間、節點／邊數、抽樣幀人工正確率等（專案未內建固定基準測試報表）。",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    doc.add_heading("與簡報草稿對齊時請修正之點", level=1)
    for line in [
        "專案名應為 child_tracker_kg，非 child_tracker_log。",
        "不存在 face_manager、alpha_tracker 模組；應使用 face_registry、yolo_tracker、video_tracker。",
        "YOLO 模型字串為 yolov8*.pt，簡報勿寫成 v11，除非實際更換模型並更新程式。",
        "無內建 interaction_timeline.png；時間軸在 interactions.json，圖表在 HTML（Chart.js）。",
        "預設輸出影片為 data/output/output.mp4，非 data/output_video.mp4。",
        "臉部主流程為 face_recognition／dlib 之距離比對策略；不宜籠統寫成「以 Cosine 為主」除非有另做實驗佐證。",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    doc.save(OUT)
    print("已寫入:", OUT)


if __name__ == "__main__":
    main()
